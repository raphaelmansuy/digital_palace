#!/usr/bin/env -S uv run
# /// script
# requires-python = ">=3.8"
# dependencies = [
#     "python-docx",
#     "Pillow",
#     "Pygments",
#     "selenium",
#     "playwright",
#     "webdriver-manager",
#     "requests"
# ]
# ///
"""
Markdown to DOCX Converter with Code Block to PNG Conversion

This script converts a markdown file to a DOCX document, with optional code block to image conversion.

Special Features:
- Mermaid diagram support: ```mermaid code blocks are converted to high-quality diagram images
- CodeSnap style: macOS window-style code blocks with syntax highlighting
- Multiple rendering engines: Playwright, Selenium, CLI tools, and online services
- Flexible image replacement: Images are always generated, but only replace text when --image-replace is used

This is a self-installing Python script using UV. No manual dependency installation required!

Usage:
    chmod +x markdown_to_docx_converter.py
    
    # Convert with auto-generated output in ./tmp directory (code blocks remain as text)
    ./markdown_to_docx_converter.py input.md
    
    # Convert and replace code blocks with images in the document
    ./markdown_to_docx_converter.py input.md --image-replace
    
    # Convert with specific output file
    ./markdown_to_docx_converter.py input.md output.docx --image-replace
    
    # Convert with custom output directory
    ./markdown_to_docx_converter.py input.md --output-dir /path/to/output
    
    # Traditional way
    uv run markdown_to_docx_converter.py input.md --image-replace

Note: Images are always generated for code blocks regardless of the --image-replace flag.
The flag only controls whether the code blocks are replaced with images in the document or kept as text.

Mermaid Setup (Optional):
    For best Mermaid diagram quality, install mermaid-cli:
    npm install -g @mermaid-js/mermaid-cli
    
    The script will automatically fall back to online rendering if CLI is unavailable.
"""

import argparse
import base64
import os
import re
import requests
import subprocess
import tempfile
import time
from pathlib import Path
from typing import List, Tuple, Dict, Any

# Required imports - UV handles these dependencies automatically
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from PIL import Image, ImageDraw, ImageFont
from pygments import highlight
from pygments.lexers import get_lexer_by_name
from pygments.formatters import ImageFormatter
from pygments.util import ClassNotFound
from pygments.lexers.special import TextLexer


class MarkdownToDocxConverter:
    def __init__(self, input_file: str, output_file: str, style: str = 'default', image_replace: bool = False):
        self.input_file = Path(input_file)
        self.output_file = Path(output_file)
        self.style = style
        self.image_replace = image_replace  # Flag to control whether to replace text with images in document
        
        # Create image directory alongside the output file
        output_path = Path(output_file)
        self.image_dir = output_path.parent
        self.base_name = output_path.stem  # filename without extension
        
        self.temp_dir = tempfile.mkdtemp(prefix="md2docx_")  # Keep for temp files
        self.image_counter = 0
        self.document = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.document.styles
        
        # Create a code style for inline code
        try:
            code_style = styles.add_style('CodeChar', WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = 'Consolas'
            code_style.font.size = Pt(10)
        except Exception:
            pass  # Style might already exist
    
    def _extract_code_blocks(self, content: str) -> Tuple[str, List[Dict[str, Any]]]:
        """Extract code blocks from markdown content and replace with placeholders"""
        code_blocks = []
        placeholder_pattern = "{{{{CODE_BLOCK_{}}}}}"  # Double braces to escape them in format()
        
        # Improved pattern to match fenced code blocks with more flexibility
        # Handles: languages with hyphens, optional newlines, various edge cases
        code_block_pattern = r'```([a-zA-Z0-9_+-]*)\n?(.*?)\n?```'
        
        def replace_code_block(match):
            language = match.group(1) if match.group(1) else 'text'
            code_content = match.group(2)
            
            block_info = {
                'language': language,
                'content': code_content,
                'index': len(code_blocks)
            }
            code_blocks.append(block_info)
            
            return placeholder_pattern.format(len(code_blocks) - 1)
        
        # Replace code blocks with placeholders
        modified_content = re.sub(code_block_pattern, replace_code_block, content, flags=re.DOTALL)
        
        return modified_content, code_blocks
    
    def _create_code_image(self, code: str, language: str = 'text') -> str:
        """Create a true retina-quality PNG image from code block (1200 DPI)"""
        self.image_counter += 1
        # Use persistent naming: outputfile_image001.png
        image_filename = f"{self.base_name}_image{self.image_counter:03d}.png"
        image_path = str(self.image_dir / image_filename)
        
        # Special handling for Mermaid diagrams
        if language.lower() == 'mermaid':
            return self._create_mermaid_image(code, image_path)
        
        # Choose styling based on user preference
        if self.style == 'codesnap':
            return self._create_codesnap_image(code, language, image_path)
        else:
            return self._create_default_image(code, language, image_path)
    
    def _create_default_image(self, code: str, language: str, image_path: str) -> str:
        """Create default style syntax highlighted image"""
        try:
            # Get lexer for the specified language
            if language and language.lower() != 'text':
                try:
                    lexer = get_lexer_by_name(language, stripall=True)
                except ClassNotFound:
                    lexer = TextLexer()
            else:
                lexer = TextLexer()
            
            # Configure the image formatter with maximum retina-quality settings
            # Scale up for highest resolution (1200 DPI equivalent)
            scale_factor = 12  # 12x scale for true retina-crisp text
            formatter = ImageFormatter(
                font_name='Monaco',
                font_size=14 * scale_factor,  # Larger font for scaling
                line_numbers=False,
                style='colorful',
                line_pad=10 * scale_factor,
                left_pad=20 * scale_factor,
                right_pad=20 * scale_factor
            )
            
            # Generate the image
            result = highlight(code, lexer, formatter)
            
            # Save the image with high DPI metadata
            with open(image_path, 'wb') as f:
                f.write(result)
            
            # Post-process to add DPI metadata for better quality
            try:
                from PIL import Image as PILImage
                img = PILImage.open(image_path)
                # Resave with true retina DPI metadata (1200 DPI)
                img.save(image_path, dpi=(1200, 1200), quality=100)
            except Exception:
                pass  # If post-processing fails, keep original
            
            return image_path
            
        except Exception as e:
            print(f"Warning: Failed to create syntax-highlighted image for {language}: {e}")
            # Fallback: create a simple text image
            return self._create_simple_text_image(code, image_path)
    
    def _create_mermaid_image(self, mermaid_code: str, image_path: str) -> str:
        """Create high-quality PNG image from Mermaid diagram code"""
        try:
            print(f"Generating Mermaid diagram: {os.path.basename(image_path)}")
            
            # Check if this is a supported diagram type
            diagram_type = self._detect_mermaid_diagram_type(mermaid_code)
            if not self._is_supported_mermaid_type(diagram_type):
                print(f"Warning: '{diagram_type}' diagrams are not supported by most Mermaid renderers")
                return self._create_mermaid_alternative_or_fallback(mermaid_code, diagram_type, image_path)
            
            # Try different Mermaid rendering methods in order of preference
            
            # Method 1: Try mermaid-cli (mmdc) if available
            try:
                return self._create_mermaid_cli_image(mermaid_code, image_path)
            except Exception as e:
                print(f"Mermaid CLI rendering failed: {e}")
            
            # Method 2: Try Playwright with Mermaid.js
            try:
                return self._create_mermaid_playwright_image(mermaid_code, image_path)
            except Exception as e:
                print(f"Playwright Mermaid rendering failed: {e}")
            
            # Method 3: Try online Mermaid service
            try:
                return self._create_mermaid_online_image(mermaid_code, image_path)
            except Exception as e:
                print(f"Online Mermaid rendering failed: {e}")
            
            # Fallback: Create text representation
            print("All Mermaid rendering methods failed, creating text fallback")
            return self._create_mermaid_text_fallback(mermaid_code, image_path)
            
        except Exception as e:
            print(f"Warning: Failed to create Mermaid image: {e}")
            return self._create_mermaid_text_fallback(mermaid_code, image_path)
    
    def _detect_mermaid_diagram_type(self, mermaid_code: str) -> str:
        """Detect the type of Mermaid diagram from the code"""
        lines = [line.strip() for line in mermaid_code.split('\n') if line.strip()]
        if not lines:
            return 'unknown'
        
        # Skip configuration lines that start with %%
        actual_content_lines = [line for line in lines if not line.startswith('%%')]
        if not actual_content_lines:
            return 'unknown'
            
        first_content_line = actual_content_lines[0].lower()
        
        # Map of diagram keywords to types
        diagram_types = {
            'flowchart': 'flowchart',
            'graph': 'flowchart',
            'sequencediagram': 'sequenceDiagram',
            'classdiagram': 'classDiagram',
            'statediagram': 'stateDiagram',
            'state-diagram': 'stateDiagram',
            'journey': 'journey',
            'gantt': 'gantt',
            'pie': 'pie',
            'gitgraph': 'gitgraph',
            'git': 'gitgraph',
            'erdiagram': 'erDiagram',
            'entityrelationshipdiagram': 'erDiagram',
            'userjourneydiagram': 'journey',
            'mindmap': 'mindmap',
            'timeline': 'timeline',
            'quadrantchart': 'quadrantChart',
            'xychart': 'xyChart',
            'xychart-beta': 'xyChart',
            'requirementdiagram': 'requirementDiagram',
            'c4context': 'c4',
            'c4container': 'c4',
            'c4component': 'c4',
            'c4dynamic': 'c4',
            'c4deployment': 'c4',
            'sankey-beta': 'sankey',
            'block-beta': 'block',
            'packet-beta': 'packet',
            'architecture-beta': 'architecture',
            'kanban': 'kanban',
            'treemap': 'treemap',
            'zenuml': 'zenuml'
        }
        
        # Check for diagram type
        for keyword, diagram_type in diagram_types.items():
            if first_content_line.startswith(keyword):
                return diagram_type
        
        # Special handling for some diagram types
        if 'gitgraph' in first_content_line or first_content_line.startswith('git'):
            return 'gitgraph'
        
        return 'unknown'
    
    def _is_supported_mermaid_type(self, diagram_type: str) -> bool:
        """Check if the diagram type is well-supported by rendering engines"""
        # Diagram types with excellent support across all rendering methods
        well_supported = {
            'flowchart', 'sequenceDiagram', 'classDiagram', 'stateDiagram',
            'gantt', 'pie', 'journey', 'erDiagram', 'mindmap', 'timeline',
            'quadrantChart', 'requirementDiagram'
        }
        
        # Diagram types with limited or problematic support (for reference)
        # limited_support = {
        #     'gitgraph',  # Limited CLI support, browser-dependent
        #     'xyChart',   # Beta feature
        #     'c4',        # Experimental
        #     'sankey',    # Beta/experimental
        #     'block',     # Beta
        #     'packet',    # Beta
        #     'architecture', # Beta
        #     'kanban',    # New/beta
        #     'treemap',   # New/beta
        #     'zenuml'     # Third-party integration
        # }
        
        return diagram_type in well_supported
    
    def _convert_gitgraph_syntax(self, mermaid_code: str) -> str:
        """Convert gitgraph/gitGraph syntax to standard git syntax if possible"""
        lines = mermaid_code.split('\n')
        converted_lines = []
        
        for line in lines:
            stripped = line.strip()
            
            # Convert gitgraph to git
            if stripped.lower().startswith('gitgraph'):
                converted_lines.append(line.replace('gitgraph', 'git', 1).replace('gitGraph', 'git', 1))
            else:
                converted_lines.append(line)
        
        return '\n'.join(converted_lines)
    
    def _create_mermaid_alternative_or_fallback(self, mermaid_code: str, diagram_type: str, image_path: str) -> str:
        """Create alternative representation for unsupported diagram types"""
        
        if diagram_type == 'gitgraph':
            # Try converting gitgraph to regular git syntax
            converted_code = self._convert_gitgraph_syntax(mermaid_code)
            if converted_code != mermaid_code:
                print("Converting gitGraph to git syntax...")
                try:
                    return self._create_mermaid_cli_image(converted_code, image_path)
                except Exception as e:
                    print(f"Converted syntax also failed: {e}")
            
            # Try alternative: create a text-based git flow representation
            print("Creating text-based git flow representation...")
            git_flow_text = self._create_git_flow_text_representation(mermaid_code)
            return self._create_simple_text_image(git_flow_text, image_path)
        
        elif diagram_type in ['c4', 'sankey', 'block', 'packet', 'architecture', 'kanban', 'treemap']:
            # For beta/experimental features, try online service first
            try:
                print(f"Trying online service for {diagram_type} diagram...")
                return self._create_mermaid_online_image(mermaid_code, image_path)
            except Exception as e:
                print(f"Online service failed for {diagram_type}: {e}")
        
        # Default fallback
        return self._create_mermaid_text_fallback(mermaid_code, image_path)
    
    def _create_git_flow_text_representation(self, mermaid_code: str) -> str:
        """Create a text representation of a git flow from mermaid gitgraph code"""
        lines = [line.strip() for line in mermaid_code.split('\n') if line.strip()]
        
        text_flow = ["GIT FLOW DIAGRAM", "=" * 50, ""]
        
        current_branch = 'main'
        branches = {'main': []}
        commit_count = 0
        
        for line in lines:
            if line.lower().startswith(('gitgraph', 'git')):
                continue
            
            if line.startswith('commit'):
                commit_count += 1
                commit_info = f"Commit {commit_count}"
                if 'id:' in line:
                    commit_id = line.split('id:')[1].strip().strip('"').strip("'")
                    commit_info = f"Commit: {commit_id}"
                if 'type:' in line:
                    commit_type = line.split('type:')[1].strip().split()[0]
                    commit_info += f" ({commit_type})"
                
                text_flow.append(f"[{current_branch}] {commit_info}")
                
            elif line.startswith('branch'):
                branch_name = line.split('branch')[1].strip()
                current_branch = branch_name
                branches[branch_name] = []
                text_flow.append("")
                text_flow.append(f"Created branch: {branch_name}")
                text_flow.append(f"Switched to: {branch_name}")
                
            elif line.startswith('checkout'):
                branch_name = line.split('checkout')[1].strip()
                current_branch = branch_name
                text_flow.append("")
                text_flow.append(f"Switched to: {branch_name}")
                
            elif line.startswith('merge'):
                merge_branch = line.split('merge')[1].strip().split()[0]
                text_flow.append("")
                text_flow.append(f"[{current_branch}] Merged branch '{merge_branch}' into '{current_branch}'")
                
            elif line.startswith('cherry-pick'):
                if 'id:' in line:
                    commit_id = line.split('id:')[1].strip().strip('"').strip("'")
                    text_flow.append(f"[{current_branch}] Cherry-picked commit: {commit_id}")
        
        text_flow.extend(["", "=" * 50, 
                         "Note: This is a text representation of the git flow.",
                         "Visual diagram could not be rendered."])
        
        return '\n'.join(text_flow)
    
    def _extract_mermaid_theme_config(self, mermaid_code: str) -> dict:
        """Extract theme configuration from Mermaid code"""
        import json
        import re
        
        # Look for %%{init: {...}}%% pattern
        theme_pattern = r'%%\{init:\s*(\{.*?\})\s*\}%%'
        matches = re.findall(theme_pattern, mermaid_code, re.DOTALL)
        
        if not matches:
            return None
        
        try:
            # Parse the JSON configuration
            config_str = matches[0]
            # Handle JavaScript-style object notation (unquoted keys)
            config_str = re.sub(r'([{,]\s*)([a-zA-Z_$][a-zA-Z0-9_$]*)\s*:', r'\1"\2":', config_str)
            config = json.loads(config_str)
            return config
        except (json.JSONDecodeError, IndexError) as e:
            print(f"Warning: Could not parse theme configuration: {e}")
            return None
    
    def _create_mermaid_config_file(self, theme_config: dict) -> str:
        """Create a mermaid configuration file for custom themes"""
        import json
        
        # Create a complete mermaid configuration
        full_config = {
            "theme": "base",
            "themeVariables": theme_config.get('themeVariables', {}),
            "flowchart": {
                "htmlLabels": True,
                "useMaxWidth": True
            },
            "sequence": {
                "diagramMarginX": 50,
                "diagramMarginY": 10,
                "actorMargin": 50,
                "width": 150,
                "height": 65,
                "boxMargin": 10,
                "boxTextMargin": 5,
                "noteMargin": 10,
                "messageMargin": 35,
                "mirrorActors": True,
                "bottomMarginAdj": 1,
                "useMaxWidth": True
            },
            "gantt": {
                "numberSectionStyles": 4,
                "axisFormat": "%m/%d/%Y",
                "useMaxWidth": True
            }
        }
        
        # Merge any additional configuration from the theme config
        if 'flowchart' in theme_config:
            full_config['flowchart'].update(theme_config['flowchart'])
        if 'sequence' in theme_config:
            full_config['sequence'].update(theme_config['sequence'])
        if 'gantt' in theme_config:
            full_config['gantt'].update(theme_config['gantt'])
        
        # Create temporary config file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as f:
            json.dump(full_config, f, indent=2)
            return f.name
    
    def _generate_mermaid_js_config(self, theme_config: dict) -> str:
        """Generate JavaScript configuration for Mermaid.js initialization"""
        import json
        
        if theme_config:
            # Use the custom theme configuration
            js_config = {
                "startOnLoad": True,
                "theme": "base",
                "themeVariables": theme_config.get('themeVariables', {}),
                "flowchart": {
                    "htmlLabels": True,
                    "useMaxWidth": True
                },
                "sequence": {
                    "useMaxWidth": True
                },
                "gantt": {
                    "useMaxWidth": True
                }
            }
            
            # Merge any additional configuration
            if 'flowchart' in theme_config:
                js_config['flowchart'].update(theme_config['flowchart'])
            if 'sequence' in theme_config:
                js_config['sequence'].update(theme_config['sequence'])
            if 'gantt' in theme_config:
                js_config['gantt'].update(theme_config['gantt'])
                
        else:
            # Default neutral theme
            js_config = {
                "startOnLoad": True,
                "theme": "neutral",
                "themeVariables": {
                    "primaryColor": "#ffffff",
                    "primaryTextColor": "#000000",
                    "primaryBorderColor": "#cccccc",
                    "lineColor": "#333333",
                    "secondaryColor": "#f8f8f8",
                    "tertiaryColor": "#e8e8e8"
                },
                "flowchart": {
                    "htmlLabels": True,
                    "useMaxWidth": True
                },
                "sequence": {
                    "useMaxWidth": True
                },
                "gantt": {
                    "useMaxWidth": True
                }
            }
        
        return json.dumps(js_config, indent=2)
    
    def _create_mermaid_cli_image(self, mermaid_code: str, image_path: str) -> str:
        """Create Mermaid image using mermaid-cli (mmdc)"""
        
        # Check if mmdc is available
        try:
            subprocess.run(['mmdc', '--version'], check=True, capture_output=True)
        except (subprocess.CalledProcessError, FileNotFoundError):
            raise Exception("mermaid-cli (mmdc) not found. Install with: npm install -g @mermaid-js/mermaid-cli")
        
        # Parse theme configuration from the mermaid code
        theme_config = self._extract_mermaid_theme_config(mermaid_code)
        
        # Create temporary mermaid file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
            f.write(mermaid_code)
            mermaid_file = f.name
        
        # Create config file if custom theme is detected
        config_file = None
        if theme_config:
            config_file = self._create_mermaid_config_file(theme_config)
        
        try:
            # Run mermaid-cli with high-quality settings
            cmd = [
                'mmdc',
                '-i', mermaid_file,
                '-o', image_path,
                '-s', '3',        # Scale factor for high resolution
                '-b', 'transparent',  # Transparent background to preserve theme colors
                '--width', '1200',    # Max width
                '--height', '800'     # Max height
            ]
            
            # Add theme configuration if available
            if config_file:
                cmd.extend(['-c', config_file])
                print("Using custom theme configuration")
            else:
                # Only use default theme if no custom theme is specified
                cmd.extend(['-t', 'neutral'])
                print("Using default neutral theme")
            
            subprocess.run(cmd, check=True, capture_output=True, text=True)
            
            # Verify the image was created
            if os.path.exists(image_path) and os.path.getsize(image_path) > 0:
                print("Successfully created Mermaid image using mermaid-cli")
                return image_path
            else:
                raise Exception("mermaid-cli produced empty or missing file")
                
        finally:
            # Cleanup temporary files
            try:
                os.unlink(mermaid_file)
                if config_file:
                    os.unlink(config_file)
            except Exception:
                pass
    
    def _create_mermaid_playwright_image(self, mermaid_code: str, image_path: str) -> str:
        """Create Mermaid image using Playwright with Mermaid.js"""
        try:
            from playwright.sync_api import sync_playwright
            
            # Extract theme configuration
            theme_config = self._extract_mermaid_theme_config(mermaid_code)
            
            # Generate JavaScript configuration for Mermaid
            mermaid_config = self._generate_mermaid_js_config(theme_config)
            
            # Create HTML file with Mermaid.js
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Mermaid Diagram</title>
    <script type="module">
        import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
        mermaid.initialize({mermaid_config});
        
        window.addEventListener('load', () => {{
            setTimeout(() => {{
                document.body.setAttribute('data-ready', 'true');
            }}, 1000);
        }});
    </script>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: white;
            margin: 20px;
            padding: 20px;
            display: flex;
            justify-content: center;
            align-items: center;
            min-height: 90vh;
        }}
        .mermaid {{
            background: white;
            border-radius: 8px;
            padding: 20px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
    </style>
</head>
<body>
    <div class="mermaid">
{mermaid_code}
    </div>
</body>
</html>"""

            # Save HTML to temporary file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                html_file = f.name

            try:
                with sync_playwright() as p:
                    browser = p.chromium.launch(headless=True)
                    page = browser.new_page(
                        device_scale_factor=3,  # High DPI for quality
                        viewport={'width': 1200, 'height': 800}
                    )
                    
                    page.goto(f'file://{html_file}')
                    
                    # Wait for Mermaid to render
                    page.wait_for_selector('[data-ready="true"]', timeout=10000)
                    page.wait_for_timeout(2000)  # Additional wait for complete rendering
                    
                    # Take screenshot of the mermaid diagram
                    mermaid_element = page.locator('.mermaid')
                    mermaid_element.screenshot(path=image_path, type='png')
                    
                    browser.close()

                # Verify and enhance image quality
                from PIL import Image as PILImage
                img = PILImage.open(image_path)
                img.save(image_path, dpi=(300, 300), quality=100, optimize=True)
                print("Successfully created Mermaid image using Playwright")
                return image_path
                    
            finally:
                try:
                    os.unlink(html_file)
                except Exception:
                    pass
                    
        except ImportError:
            raise Exception("Playwright not available for Mermaid rendering")
    
    def _create_mermaid_online_image(self, mermaid_code: str, image_path: str) -> str:
        """Create Mermaid image using online Mermaid service"""
        
        try:
            # Use Mermaid.ink service (public API)
            # Encode the mermaid code for URL
            encoded_diagram = base64.b64encode(mermaid_code.encode('utf-8')).decode('ascii')
            
            # Mermaid.ink API endpoint
            url = f"https://mermaid.ink/img/{encoded_diagram}"
            
            # Add parameters for better quality
            params = {
                'type': 'png',
                'theme': 'neutral',
                'scale': '3'  # Higher scale for better quality
            }
            
            # Make request with longer timeout for complex diagrams
            response = requests.get(url, params=params, timeout=30)
            response.raise_for_status()
            
            # Check if we got a valid image
            if response.headers.get('content-type', '').startswith('image/'):
                with open(image_path, 'wb') as f:
                    f.write(response.content)
                
                # Verify the image was created and is valid
                if os.path.getsize(image_path) > 1000:  # Reasonable minimum size
                    from PIL import Image as PILImage
                    try:
                        # Verify it's a valid image and enhance DPI
                        img = PILImage.open(image_path)
                        img.save(image_path, dpi=(300, 300), quality=100, optimize=True)
                        print("Successfully created Mermaid image using online service")
                        return image_path
                    except Exception:
                        raise Exception("Online service returned invalid image data")
                else:
                    raise Exception("Online service returned empty or tiny image")
            else:
                raise Exception(f"Online service returned non-image content: {response.headers.get('content-type')}")
                
        except requests.RequestException as e:
            raise Exception(f"Online Mermaid service request failed: {e}")
        except Exception as e:
            raise Exception(f"Online Mermaid service error: {e}")
    
    def _create_mermaid_text_fallback(self, mermaid_code: str, image_path: str) -> str:
        """Create text-based fallback for Mermaid diagrams"""
        try:
            # Create a text representation of the Mermaid diagram
            fallback_text = f"[MERMAID DIAGRAM]\n\n{mermaid_code}\n\n[Diagram could not be rendered as image]"
            
            return self._create_simple_text_image(fallback_text, image_path)
            
        except Exception as e:
            print(f"Error creating Mermaid text fallback: {e}")
            return None
    
    def _create_codesnap_image(self, code: str, language: str, image_path: str) -> str:
        """Create CodeSnap style image with macOS window frame using headless browser rendering"""
        try:
            import tempfile
            import os
            
            # Get lexer for syntax highlighting
            if language and language.lower() != 'text':
                try:
                    lexer = get_lexer_by_name(language, stripall=True)
                except ClassNotFound:
                    lexer = TextLexer()
            else:
                lexer = TextLexer()
            
            # Generate syntax-highlighted HTML using Pygments
            from pygments.formatters import HtmlFormatter
            
            # Create a custom HTML formatter with high-quality settings
            formatter = HtmlFormatter(
                style='monokai',  # Dark theme that works well with codesnap
                noclasses=True,   # Inline styles for better control
                linenos=False,
                cssclass='highlight',
                prestyles='margin: 0; padding: 0; font-family: Monaco, Menlo, "Ubuntu Mono", monospace;'
            )
            
            # Generate highlighted code HTML
            highlighted_code = highlight(code, lexer, formatter)
            
            
            # Create complete HTML5 document optimized for headless browser rendering
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Code Snapshot</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: linear-gradient(135deg, #1e1e1e 0%, #2d2d2d 100%);
            padding: 40px;
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            font-smooth: always;
            -webkit-font-smoothing: antialiased;
            -moz-osx-font-smoothing: grayscale;
        }}
        
        .window {{
            background: #2d2d2d;
            border-radius: 12px;
            box-shadow: 
                0 32px 64px rgba(0, 0, 0, 0.35),
                0 8px 16px rgba(0, 0, 0, 0.2),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
            overflow: hidden;
            min-width: 600px;
            max-width: 1000px;
            width: fit-content;
            backdrop-filter: blur(20px);
        }}
        
        .title-bar {{
            background: linear-gradient(to bottom, 
                rgba(64, 64, 64, 0.95) 0%, 
                rgba(56, 56, 56, 0.95) 100%);
            height: 44px;
            display: flex;
            align-items: center;
            padding: 0 20px;
            border-bottom: 1px solid rgba(42, 42, 42, 0.8);
            backdrop-filter: blur(10px);
        }}
        
        .traffic-lights {{
            display: flex;
            gap: 8px;
            align-items: center;
        }}
        
        .dot {{
            width: 12px;
            height: 12px;
            border-radius: 50%;
            border: 0.5px solid rgba(0, 0, 0, 0.15);
            box-shadow: 
                inset 0 1px 0 rgba(255, 255, 255, 0.2),
                0 1px 1px rgba(0, 0, 0, 0.1);
        }}
        
        .dot.close {{ 
            background: linear-gradient(135deg, #ff6058 0%, #ff4d4d 100%);
        }}
        .dot.minimize {{ 
            background: linear-gradient(135deg, #ffbd30 0%, #ffab00 100%);
        }}
        .dot.maximize {{ 
            background: linear-gradient(135deg, #28ca44 0%, #20a034 100%);
        }}
        
        .title {{
            flex: 1;
            text-align: center;
            color: #cccccc;
            font-size: 13px;
            font-weight: 500;
            text-shadow: 0 1px 0 rgba(0, 0, 0, 0.5);
        }}
        
        .content {{
            background: linear-gradient(to bottom, #282c34 0%, #21252b 100%);
            padding: 28px;
            font-family: 'SF Mono', Monaco, Menlo, 'Ubuntu Mono', Consolas, monospace;
            font-size: 14px;
            line-height: 1.7;
            color: #abb2bf;
            overflow-x: auto;
            border-radius: 0 0 12px 12px;
        }}
        
        /* Enhanced Pygments styles for better codesnap appearance */
        .highlight {{
            background: transparent !important;
            font-family: 'SF Mono', Monaco, Menlo, 'Ubuntu Mono', Consolas, monospace !important;
            font-size: 14px !important;
            line-height: 1.7 !important;
            font-weight: 400 !important;
        }}
        
        .highlight pre {{
            margin: 0 !important;
            padding: 0 !important;
            background: transparent !important;
            white-space: pre-wrap;
            word-wrap: break-word;
            font-feature-settings: "liga" 0, "calt" 0;
        }}
        
        /* Ensure crisp text rendering for screenshots */
        pre, code {{
            font-feature-settings: "liga" 0, "calt" 0;
            -webkit-font-smoothing: antialiased;
            -moz-osx-font-smoothing: grayscale;
            text-rendering: optimizeSpeed;
        }}
        
        /* Force hardware acceleration for smooth rendering */
        .window {{
            transform: translateZ(0);
            will-change: transform;
        }}
    </style>
</head>
<body>
    <div class="window">
        <div class="title-bar">
            <div class="traffic-lights">
                <div class="dot close"></div>
                <div class="dot minimize"></div>
                <div class="dot maximize"></div>
            </div>
            <div class="title">{language if language else 'code'}</div>
            <div style="width: 60px;"></div> <!-- Spacer for centering -->
        </div>
        <div class="content">
            {highlighted_code}
        </div>
    </div>
</body>
</html>"""
            
            # Save HTML to temporary file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                html_file = f.name
            
            try:
                # Priority 1: Try headless browser rendering (Playwright)
                try:
                    return self._create_playwright_screenshot(html_file, image_path)
                except Exception as e:
                    print(f"Playwright rendering failed: {e}")
                
                # Priority 2: Try Selenium headless browser
                try:
                    print("Attempting Selenium rendering...")
                    return self._create_selenium_screenshot(html_file, image_path)
                except Exception as e:
                    print(f"Selenium rendering failed: {e}")
                
                # Priority 3: Try wkhtmltopdf (if available)
                try:
                    print("Attempting wkhtmltopdf rendering...")
                    return self._create_wkhtmltopdf_image(html_file, image_path)
                except Exception as e:
                    print(f"wkhtmltopdf rendering failed: {e}")
                
                # Priority 4: Enhanced PIL fallback
                print("Using enhanced PIL fallback...")
                return self._create_codesnap_pil_fallback(code, language, image_path)
                
            finally:
                # Cleanup HTML file
                try:
                    os.unlink(html_file)
                except Exception:
                    pass
            
        except Exception as e:
            print(f"Warning: Failed to create HTML5 CodeSnap image for {language}: {e}")
            # Final fallback to PIL-based rendering
            return self._create_codesnap_pil_fallback(code, language, image_path)
    
    def _create_playwright_screenshot(self, html_file: str, image_path: str) -> str:
        """Create high-quality screenshot using Playwright headless browser"""
        try:
            from playwright.sync_api import sync_playwright
            from PIL import Image as PILImageModule
            
            with sync_playwright() as p:
                # Launch Chromium with maximum DPI settings for true retina quality
                browser = p.chromium.launch(
                    headless=True,
                    args=[
                        '--force-device-scale-factor=12',  # 12x scale for true retina quality
                        '--high-dpi-support=1',
                        '--force-color-profile=srgb',
                        '--disable-background-timer-throttling',
                        '--disable-renderer-backgrounding',
                        '--disable-backgrounding-occluded-windows',
                        '--font-render-hinting=none',  # Better font rendering
                        '--enable-font-antialiasing',
                        '--disable-lcd-text',  # Force grayscale antialiasing for crisp text
                        '--force-prefers-reduced-motion',
                        '--disable-background-media-processing',
                        '--disable-renderer-accessibility',
                    ]
                )
                
                # Create page with maximum device scale factor for true retina
                page = browser.new_page(
                    device_scale_factor=12,  # 12x for true retina quality (1200 DPI equivalent)
                    viewport={'width': 2000, 'height': 1600}  # Larger viewport for better text rendering
                )
                
                # Navigate to HTML file
                page.goto(f'file://{html_file}')
                
                # Wait for rendering to complete
                page.wait_for_load_state('networkidle')
                page.wait_for_timeout(1000)  # Extended wait for font loading and rendering
                
                # Ensure the window element exists and is visible
                page.wait_for_selector('.window', state='visible')
                
                # Get the window element for precise cropping
                window_element = page.locator('.window')
                
                # Take screenshot with high quality settings
                window_element.screenshot(
                    path=image_path,
                    type='png',
                    quality=100,
                    animations='disabled'
                )
                
                browser.close()
                
                # Post-process to add maximum DPI metadata for true retina
                img = PILImageModule.open(image_path)
                img.save(image_path, dpi=(1200, 1200), quality=100, optimize=True)  # 1200 DPI for true retina quality
                
                print("Successfully created CodeSnap image using Playwright")
                return image_path
                
        except ImportError:
            raise Exception("Playwright not available - install with: pip install playwright && playwright install chromium")
        except Exception as e:
            raise Exception(f"Playwright screenshot failed: {e}")
    
    def _create_selenium_screenshot(self, html_file: str, image_path: str) -> str:
        """Create screenshot using Selenium headless browser"""
        try:
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
            from selenium.webdriver.common.by import By
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            from PIL import Image as PILImageModule
            
            # Configure Chrome options for ultra-high-quality rendering
            chrome_options = Options()
            chrome_options.add_argument('--headless=new')  # Use new headless mode
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--force-device-scale-factor=6')  # 6x scale for optimal retina quality
            chrome_options.add_argument('--high-dpi-support=1')
            chrome_options.add_argument('--force-color-profile=srgb')
            chrome_options.add_argument('--window-size=1800,1400')  # Optimal window for better text rendering
            chrome_options.add_argument('--disable-background-timer-throttling')
            chrome_options.add_argument('--disable-renderer-backgrounding')
            chrome_options.add_argument('--disable-backgrounding-occluded-windows')
            chrome_options.add_argument('--font-render-hinting=none')  # Better font rendering
            chrome_options.add_argument('--enable-font-antialiasing')
            chrome_options.add_argument('--disable-lcd-text')  # Force grayscale antialiasing
            chrome_options.add_argument('--disable-web-security')  # Allow local file access
            chrome_options.add_argument('--allow-running-insecure-content')
            chrome_options.add_argument('--disable-features=VizDisplayCompositor')
            
            # Launch browser with ChromeDriverManager for automatic driver management
            try:
                from selenium.webdriver.chrome.service import Service
                from webdriver_manager.chrome import ChromeDriverManager
                service = Service(ChromeDriverManager().install())
                driver = webdriver.Chrome(service=service, options=chrome_options)
            except ImportError:
                # Fallback to regular Chrome driver
                driver = webdriver.Chrome(options=chrome_options)
            
            try:
                # Navigate to HTML file
                driver.get(f'file://{html_file}')
                
                # Wait for the window element to be present and fully rendered
                wait = WebDriverWait(driver, 10)
                window_element = wait.until(
                    EC.presence_of_element_located((By.CLASS_NAME, "window"))
                )
                
                # Extended wait for font loading and rendering
                driver.implicitly_wait(2)
                time.sleep(1)  # Additional wait to ensure complete rendering
                
                # Take screenshot of just the window element
                print(f"Taking screenshot with Selenium at {chrome_options.arguments}")
                window_element.screenshot(image_path)
                
                # Verify the screenshot was taken properly
                if os.path.getsize(image_path) < 500:  # Very small file suggests blank image
                    print("Warning: Screenshot appears to be blank, taking full page screenshot as fallback")
                    driver.save_screenshot(image_path.replace('.png', '_fullpage.png'))
                    # Try to crop to window area
                    from PIL import Image as PILImage
                    full_img = PILImage.open(image_path.replace('.png', '_fullpage.png'))
                    # Save as the main image
                    full_img.save(image_path)
                    os.remove(image_path.replace('.png', '_fullpage.png'))
                
                # Post-process to add maximum DPI metadata for true retina
                img = PILImageModule.open(image_path)
                img.save(image_path, dpi=(1200, 1200), quality=100, optimize=True)  # 1200 DPI for true retina quality
                
                print("Successfully created CodeSnap image using Selenium")
                return image_path
                
            finally:
                driver.quit()
                
        except ImportError:
            raise Exception("Selenium not available - install with: pip install selenium")
        except Exception as e:
            raise Exception(f"Selenium screenshot failed: {e}")
    
    def _create_wkhtmltopdf_image(self, html_file: str, image_path: str) -> str:
        """Create image using wkhtmltopdf (fallback method)"""
        import tempfile
        import subprocess
        
        # Create temporary PDF
        temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_pdf.close()
        
        try:
            # wkhtmltopdf command with high DPI settings
            cmd = [
                'wkhtmltopdf',
                '--page-size', 'A4',
                '--orientation', 'Portrait', 
                '--margin-top', '0',
                '--margin-right', '0',
                '--margin-bottom', '0',
                '--margin-left', '0',
                '--dpi', '300',  # 300 DPI for high quality
                '--image-quality', '100',
                '--disable-smart-shrinking',
                '--zoom', '2.0',  # 2x zoom for better quality
                '--javascript-delay', '1000',  # Wait for rendering
                html_file,
                temp_pdf.name
            ]
            
            # Check if wkhtmltopdf is available and run
            subprocess.run(['which', 'wkhtmltopdf'], check=True, capture_output=True)
            subprocess.run(cmd, check=True, capture_output=True)
            
            # Convert PDF to PNG using ImageMagick or PIL
            try:
                # Try ImageMagick first for best quality
                convert_cmd = [
                    'convert',
                    '-density', '300',  # 300 DPI
                    '-quality', '100',
                    '-background', 'transparent',
                    temp_pdf.name,
                    image_path
                ]
                subprocess.run(convert_cmd, check=True, capture_output=True)
                
            except (subprocess.CalledProcessError, FileNotFoundError):
                # Fallback: use pdf2image if available
                try:
                    from pdf2image import convert_from_path
                    images = convert_from_path(temp_pdf.name, dpi=300, first_page=1, last_page=1)
                    if images:
                        images[0].save(image_path, 'PNG', dpi=(300, 300), quality=100)
                except ImportError:
                    raise Exception("Neither ImageMagick nor pdf2image available for PDF conversion")
            
            print("Successfully created CodeSnap image using wkhtmltopdf")
            return image_path
            
        finally:
            # Cleanup
            try:
                os.unlink(temp_pdf.name)
            except Exception:
                pass
    
    def _create_codesnap_pil_fallback(self, code: str, language: str, image_path: str) -> str:
        """Enhanced PIL-based fallback with better syntax highlighting and 800 DPI ultra-high quality"""
        try:
            from PIL import Image as PILImage, ImageDraw, ImageFont
            
            # Get lexer for syntax highlighting
            if language and language.lower() != 'text':
                try:
                    lexer = get_lexer_by_name(language, stripall=True)
                except ClassNotFound:
                    lexer = TextLexer()
            else:
                lexer = TextLexer()
            
            # 1200 DPI settings - scale factor for true retina-quality rendering
            dpi_scale = 15  # 15x scale for 1200 DPI true retina quality
            window_border = 8 * dpi_scale
            title_bar_height = 44 * dpi_scale
            dot_size = 12 * dpi_scale
            dot_spacing = 20 * dpi_scale
            content_padding = 24 * dpi_scale
            font_size = 14 * dpi_scale
            line_height = 22 * dpi_scale
            
            # Enhanced window colors (more accurate macOS style)
            window_bg = (45, 45, 45)      # Darker gray
            title_bar_bg = (64, 64, 64)   # Title bar color
            content_bg = (40, 44, 52)     # Code background (VS Code dark)
            text_color = (171, 178, 191)  # Default text color
            dot_colors = [
                (255, 95, 87),   # Red (close)
                (255, 189, 46),  # Yellow (minimize)
                (40, 202, 66)    # Green (maximize)
            ]
            
            # Calculate text dimensions with proper font metrics
            lines = code.split('\n')
            max_line_length = max(len(line.expandtabs(4)) for line in lines) if lines else 50
            
            # Load high-quality font
            try:
                font = ImageFont.truetype("/System/Library/Fonts/SF-Mono-Regular.otf", font_size)
            except Exception:
                try:
                    font = ImageFont.truetype("/System/Library/Fonts/Monaco.ttc", font_size)
                except Exception:
                    try:
                        font = ImageFont.truetype("/System/Library/Fonts/Menlo.ttc", font_size)
                    except Exception:
                        font = ImageFont.load_default()
            
            # Calculate precise image dimensions
            # Get actual text width using font metrics
            try:
                # Use the font to get actual character width
                sample_text = "M" * max_line_length
                bbox = font.getbbox(sample_text)
                content_width = bbox[2] + 2 * content_padding
            except Exception:
                # Fallback calculation
                content_width = max_line_length * (font_size * 0.6) + 2 * content_padding
            
            content_height = len(lines) * line_height + 2 * content_padding
            
            total_width = int(content_width + 2 * window_border)
            total_height = int(content_height + title_bar_height + 2 * window_border)
            
            # Create high-resolution image
            img = PILImage.new('RGB', (total_width, total_height), window_bg)
            draw = ImageDraw.Draw(img)
            
            # Draw title bar with gradient effect (simulated)
            for i in range(int(title_bar_height)):
                alpha = i / title_bar_height
                color_r = int(title_bar_bg[0] * (1 - alpha * 0.1))
                color_g = int(title_bar_bg[1] * (1 - alpha * 0.1))
                color_b = int(title_bar_bg[2] * (1 - alpha * 0.1))
                draw.line([
                    (window_border, window_border + i),
                    (total_width - window_border, window_border + i)
                ], fill=(color_r, color_g, color_b))
            
            # Draw window dots (traffic lights) with subtle shadows
            dot_y = window_border + title_bar_height // 2
            dot_start_x = window_border + dot_spacing
            
            for i, color in enumerate(dot_colors):
                dot_x = dot_start_x + i * dot_spacing
                # Draw subtle shadow
                shadow_offset = dpi_scale
                draw.ellipse([
                    dot_x - dot_size//2 + shadow_offset, 
                    dot_y - dot_size//2 + shadow_offset,
                    dot_x + dot_size//2 + shadow_offset, 
                    dot_y + dot_size//2 + shadow_offset
                ], fill=(0, 0, 0, 30))
                # Draw main dot
                draw.ellipse([
                    dot_x - dot_size//2, 
                    dot_y - dot_size//2,
                    dot_x + dot_size//2, 
                    dot_y + dot_size//2
                ], fill=color)
            
            # Draw title text
            if language:
                title_font_size = int(13 * dpi_scale)
                try:
                    title_font = ImageFont.truetype("/System/Library/Fonts/SF-Pro-Display-Medium.otf", title_font_size)
                except Exception:
                    title_font = font
                
                title_bbox = title_font.getbbox(language)
                title_x = (total_width - title_bbox[2]) // 2
                title_y = window_border + (title_bar_height - title_bbox[3]) // 2
                draw.text((title_x, title_y), language, fill=(204, 204, 204), font=title_font)
            
            # Draw content area
            content_y_start = window_border + title_bar_height
            draw.rectangle([
                window_border,
                content_y_start,
                total_width - window_border,
                total_height - window_border
            ], fill=content_bg)
            
            # Enhanced syntax highlighting using Pygments tokens
            try:
                from pygments.token import Token
                tokens = list(lexer.get_tokens(code))
                
                # Define color scheme (Monokai-inspired)
                token_colors = {
                    Token.Comment: (117, 113, 94),           # Comments
                    Token.Comment.Single: (117, 113, 94),
                    Token.Comment.Multiline: (117, 113, 94),
                    Token.Keyword: (249, 38, 114),           # Keywords (def, class, etc.)
                    Token.Keyword.Namespace: (249, 38, 114),
                    Token.Keyword.Type: (102, 217, 239),
                    Token.String: (230, 219, 116),           # Strings
                    Token.String.Double: (230, 219, 116),
                    Token.String.Single: (230, 219, 116),
                    Token.Number: (174, 129, 255),           # Numbers
                    Token.Number.Integer: (174, 129, 255),
                    Token.Number.Float: (174, 129, 255),
                    Token.Name.Function: (166, 226, 46),     # Function names
                    Token.Name.Class: (166, 226, 46),        # Class names
                    Token.Name.Builtin: (102, 217, 239),     # Built-ins
                    Token.Operator: (249, 38, 114),          # Operators
                    Token.Punctuation: text_color,           # Punctuation
                }
                
                # Render tokens with syntax highlighting
                x_pos = window_border + content_padding
                y_pos = content_y_start + content_padding
                
                for token_type, text in tokens:
                    if text == '\n':
                        y_pos += line_height
                        x_pos = window_border + content_padding
                        continue
                    
                    # Get color for token type
                    color = text_color  # Default
                    for token_key, token_color in token_colors.items():
                        if token_type in token_key:
                            color = token_color
                            break
                    
                    # Draw text
                    draw.text((x_pos, y_pos), text, fill=color, font=font)
                    
                    # Update x position
                    try:
                        bbox = font.getbbox(text)
                        x_pos += bbox[2]
                    except Exception:
                        x_pos += len(text) * (font_size * 0.6)
                
            except Exception:
                # Fallback to simple syntax highlighting
                y_pos = content_y_start + content_padding
                for line in lines:
                    # Simple syntax highlighting
                    line_expanded = line.expandtabs(4)
                    line_color = text_color
                    
                    # Basic highlighting rules
                    stripped = line.strip()
                    if stripped.startswith('#'):
                        line_color = (117, 113, 94)  # Comments
                    elif any(keyword in line for keyword in ['def ', 'class ', 'import ', 'from ', 'if ', 'for ', 'while ', 'try:', 'except:']):
                        line_color = (249, 38, 114)  # Keywords
                    elif '"' in line or "'" in line:
                        line_color = (230, 219, 116)  # Strings
                    
                    draw.text(
                        (window_border + content_padding, y_pos),
                        line_expanded,
                        fill=line_color,
                        font=font
                    )
                    y_pos += line_height
            
            # Save with true retina DPI metadata (1200 DPI)
            img.save(image_path, dpi=(1200, 1200), quality=100, optimize=True)
            return image_path
            
        except Exception as e:
            print(f"Warning: Enhanced PIL fallback failed: {e}")
            # Final fallback to original simple method
            return self._create_default_image(code, language, image_path)
    
    def _create_simple_text_image(self, code: str, image_path: str) -> str:
        """Create a simple text image as fallback with ultra-high DPI (800 DPI)"""
        try:
            # Calculate image size based on text with high DPI scaling
            lines = code.split('\n')
            max_line_length = max(len(line) for line in lines) if lines else 50
            
            # High DPI settings (1200 DPI equivalent for true retina)
            dpi_scale = 15  # 15x scale for 1200 DPI true retina quality
            char_width = 9 * dpi_scale
            char_height = 18 * dpi_scale
            padding = 25 * dpi_scale
            
            width = max(500 * dpi_scale, max_line_length * char_width + padding * 2)
            height = max(120 * dpi_scale, len(lines) * char_height + padding * 2)
            
            # Create high-resolution image
            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)
            
            # Try to use a high-quality monospace font
            font_size = 16 * dpi_scale
            try:
                font = ImageFont.truetype('/System/Library/Fonts/SF-Mono-Regular.otf', font_size)
            except Exception:
                try:
                    font = ImageFont.truetype('/System/Library/Fonts/Monaco.ttc', font_size)
                except Exception:
                    try:
                        font = ImageFont.truetype('Monaco', font_size)
                    except Exception:
                        try:
                            font = ImageFont.truetype('/System/Library/Fonts/Menlo.ttc', font_size)
                        except Exception:
                            # Use default but scale it
                            font = ImageFont.load_default()
            
            # Draw text with high quality
            y_offset = padding
            for line in lines:
                draw.text((padding, y_offset), line, fill='black', font=font)
                y_offset += char_height
            
            # Add border
            border_width = 2 * dpi_scale
            draw.rectangle([0, 0, width-border_width, height-border_width], 
                         outline='gray', width=border_width)
            
            # Save with true retina DPI metadata (1200 DPI)
            img.save(image_path, dpi=(1200, 1200), quality=100)
            return image_path
            
        except Exception as e:
            print(f"Error creating fallback image: {e}")
            return None
    
    def _process_markdown_element(self, line: str):
        """Process different markdown elements"""
        line = line.strip()
        
        if not line:
            # Empty line
            self.document.add_paragraph()
            return
        
        # Check for code block placeholder
        placeholder_match = re.match(r'\{\{CODE_BLOCK_(\d+)\}\}', line)
        if placeholder_match:
            # This will be handled separately
            return
        
        # Headers
        if line.startswith('#'):
            level = 0
            while level < len(line) and line[level] == '#':
                level += 1
            
            if level <= 6:
                header_text = line[level:].strip()
                if level == 1:
                    heading = self.document.add_heading(header_text, level=1)
                elif level == 2:
                    heading = self.document.add_heading(header_text, level=2)
                elif level == 3:
                    heading = self.document.add_heading(header_text, level=3)
                else:
                    heading = self.document.add_heading(header_text, level=4)
                
                # Ensure heading text is black
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
                return
        
        # Horizontal rule
        if line.strip() == '---':
            # Add spacing before the horizontal rule
            self.document.add_paragraph()
            # Create a proper horizontal line using a paragraph with bottom border
            p = self.document.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            # Get the paragraph format
            pPr = p._element.get_or_add_pPr()
            # Create border element
            pBdr = OxmlElement('w:pBdr')
            # Add bottom border
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '24')  # Thicker border size
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), '000000')  # Black color
            pBdr.append(bottom)
            pPr.append(pBdr)
            # Add spacing after the horizontal rule
            self.document.add_paragraph()
            return
        
        # Lists
        if re.match(r'^[\s]*[-\*\+]\s+', line):
            # Bullet list
            text = re.sub(r'^[\s]*[-\*\+]\s+', '', line)
            p = self.document.add_paragraph(text, style='List Bullet')
            self._apply_text_formatting(p, text)
            return
        
        if re.match(r'^[\s]*\d+\.\s+', line):
            # Numbered list
            text = re.sub(r'^[\s]*\d+\.\s+', '', line)
            p = self.document.add_paragraph(text, style='List Number')
            self._apply_text_formatting(p, text)
            return
        
        # Regular paragraph
        p = self.document.add_paragraph()
        self._apply_text_formatting(p, line)
    
    def _apply_text_formatting(self, paragraph, text: str):
        """Apply bold, italic, and inline code formatting"""
        # Clear existing text
        paragraph.clear()
        
        # Improved pattern to match formatting - order matters!
        # 1. **bold** (must come before *italic* to avoid conflicts)
        # 2. *italic* 
        # 3. `code`
        # Using non-greedy matching and proper escaping
        pattern = r'(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)'
        parts = re.split(pattern, text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # Bold text - ensure it's actually bold formatting
                content = part[2:-2]
                run = paragraph.add_run(content)
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and len(part) > 2 and not part.startswith('**'):
                # Italic text - ensure it's not bold and has content
                content = part[1:-1]
                run = paragraph.add_run(content)
                run.italic = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                # Inline code - ensure it has content
                content = part[1:-1]
                run = paragraph.add_run(content)
                try:
                    run.style = 'CodeChar'
                except Exception:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
            else:
                # Regular text
                paragraph.add_run(part)
    
    def convert(self):
        """Convert the markdown file to DOCX"""
        try:
            # Read the markdown file
            with open(self.input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extract code blocks
            modified_content, code_blocks = self._extract_code_blocks(content)
            
            # Always create images for code blocks (regardless of image_replace flag)
            code_images = {}
            for i, block in enumerate(code_blocks):
                image_path = self._create_code_image(block['content'], block['language'])
                if image_path:
                    code_images[i] = image_path
            
            # Process markdown content line by line
            lines = modified_content.split('\n')
            i = 0
            while i < len(lines):
                line = lines[i]
                
                # Check for code block placeholder
                placeholder_match = re.match(r'\{\{CODE_BLOCK_(\d+)\}\}', line.strip())
                if placeholder_match:
                    block_index = int(placeholder_match.group(1))
                    
                    if self.image_replace and block_index in code_images:
                        # Replace with image only if --image-replace flag is set
                        paragraph = self.document.add_paragraph()
                        run = paragraph.add_run()
                        try:
                            run.add_picture(code_images[block_index], width=Inches(6))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print(f"Warning: Could not add image for code block {block_index}: {e}")
                            # Fallback: add as text
                            self._add_code_block_as_text(code_blocks[block_index])
                    else:
                        # Keep as text in document (default behavior when --image-replace is not set)
                        self._add_code_block_as_text(code_blocks[block_index])
                else:
                    self._process_markdown_element(line)
                
                i += 1
            
            # Save the document
            self.document.save(self.output_file)
            print(f"Successfully converted {self.input_file} to {self.output_file}")
            if self.image_counter > 0:
                image_status = "Generated" if not self.image_replace else "Generated and embedded"
                print(f"{image_status} {self.image_counter} code block images: {self.base_name}_image001.png to {self.base_name}_image{self.image_counter:03d}.png")
                if not self.image_replace:
                    print("Note: Code blocks remain as text in document. Use --image-replace to embed images.")
            
        except Exception as e:
            print(f"Error during conversion: {e}")
            raise
        finally:
            # Cleanup temporary files
            self._cleanup()
    
    def _add_code_block_as_text(self, code_block):
        """Add code block as formatted text in the document"""
        # Add language label if available
        if code_block['language']:
            lang_paragraph = self.document.add_paragraph()
            lang_run = lang_paragraph.add_run(f"[{code_block['language']}]")
            lang_run.font.name = 'Consolas'
            lang_run.font.size = Pt(9)
            lang_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        
        # Add code content
        code_paragraph = self.document.add_paragraph()
        code_run = code_paragraph.add_run(code_block['content'])
        try:
            code_run.style = 'CodeChar'
        except Exception:
            code_run.font.name = 'Consolas'
            code_run.font.size = Pt(10)
            code_run.font.color.rgb = RGBColor(64, 64, 64)  # Dark gray for code
    
    def _cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Warning: Could not clean up temporary directory: {e}")


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to DOCX with code blocks as images'
    )
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('output', nargs='?', default=None, help='Output DOCX file (optional if --output-dir is used)')
    parser.add_argument('--output-dir', '-o', 
                       help='Output directory (defaults to ./tmp relative to input file)')
    parser.add_argument('--style', choices=['default', 'codesnap'], default='codesnap',
                       help='Code block image style: default or codesnap (macOS window style)')
    parser.add_argument('--image-replace', action='store_true',
                       help='Replace code blocks with images in the document (images are always generated)')
    parser.add_argument('--verbose', '-v', action='store_true', 
                       help='Verbose output')
    
    args = parser.parse_args()
    
    # Check if input file exists
    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' does not exist")
        return 1
    
    # Determine output file path
    input_path = Path(args.input)
    
    if args.output:
        # If output file is explicitly provided, use it
        output_file = args.output
        output_dir = os.path.dirname(output_file)
    elif args.output_dir:
        # If output-dir is provided, use it for output location
        output_dir = args.output_dir
        output_filename = input_path.stem + '.docx'
        output_file = os.path.join(output_dir, output_filename)
    else:
        # Default: create ./tmp directory alongside input file
        default_output_dir = input_path.parent / 'tmp'
        output_dir = str(default_output_dir)
        output_filename = input_path.stem + '.docx'
        output_file = str(default_output_dir / output_filename)
    
    # Create output directory if it doesn't exist
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        if args.verbose:
            print(f"Created output directory: {output_dir}")
    
    if args.verbose:
        print(f"Converting: {args.input} -> {output_file}")
        if args.image_replace:
            print("Mode: Code blocks will be replaced with images in the document")
        else:
            print("Mode: Code blocks will remain as text in document (images generated separately)")
    
    # Convert the file
    try:
        converter = MarkdownToDocxConverter(args.input, output_file, args.style, args.image_replace)
        converter.convert()
        return 0
    except Exception as e:
        print(f"Conversion failed: {e}")
        return 1


if __name__ == '__main__':
    # This script is now self-installing thanks to UV!
    # Usage examples:
    # ./markdown_to_docx_converter.py input.md                     # -> ./tmp/input.docx (text), images generated separately
    # ./markdown_to_docx_converter.py input.md --image-replace     # -> ./tmp/input.docx (with images embedded)
    # ./markdown_to_docx_converter.py input.md output.docx         # -> output.docx (text), images generated
    # ./markdown_to_docx_converter.py input.md -o /custom/dir      # -> /custom/dir/input.docx (text)
    # ./markdown_to_docx_converter.py input.md --image-replace -o /custom/dir  # -> /custom/dir/input.docx (with images)
    # UV will automatically handle all dependency installation and environment setup
    exit(main())
